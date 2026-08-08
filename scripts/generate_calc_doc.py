#!/usr/bin/env python
"""מייצר מסמך Word עם חישובי מס, NOPAT, IC_RAW, ROIC, EV/FCF וחוב."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = ROOT / "תיעוד_חישובים.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x60, 0x60, 0x60)

TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_font(run, name: str = "Calibri", size: float = 11, *,
             color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), name)


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.insert(0, bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH:
        raise ValueError("רוחב עמודות הטבלה חייב להסתכם ל-9360 DXA")

    table.autofit = False
    tbl_pr = table._tbl.tblPr

    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.insert(0, bidi)

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc: Document, text: str, *, color: RGBColor | None = None):
    paragraph = doc.add_paragraph()
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    set_font(run, color=color)
    return paragraph


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shading)
    run = paragraph.add_run(text.strip())
    set_font(run, "Consolas", 9.5)


def add_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("עמודה", "סוג", "נוסחה / משמעות")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        set_rtl(paragraph)
        run = paragraph.add_run(text)
        set_font(run, size=10, bold=True)
        shade_cell(cell, "E8EEF5")
    set_repeat_header(table.rows[0])

    for values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, values)):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            set_rtl(paragraph)
            run = paragraph.add_run(text)
            set_font(run, "Consolas" if index == 0 else "Calibri",
                     8.5 if index == 0 else 9.5)

    set_table_geometry(table, [3900, 1100, 4360])


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("עמוד ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        normal_fonts.set(qn(key), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style_fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("w:ascii", "w:hAnsi", "w:cs"):
            style_fonts.set(qn(key), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    set_rtl(header)
    set_font(header.add_run("מדריך חישובי ROIC, EV/FCF וחוב · GuruFocus"), size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    set_rtl(footer)
    add_page_field(footer)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    set_rtl(paragraph)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(
        title.add_run("חישובים רבעוניים: NOPAT, IC_RAW, ROIC, EV/FCF וחוב"),
        size=28,
        color=DARK_BLUE,
        bold=True,
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(
        subtitle.add_run("מפרט הנוסחאות המחושבות בפרויקט"),
        size=13.5,
        color=MUTED,
    )

    add_paragraph(
        doc,
        "המסמך כולל רק את החישובים הפעילים בפלט. RawTaxRate משמש ישירות "
        "בחישוב NOPAT.",
    )

    add_heading(doc, "1. RawTaxRate ו-NOPAT רבעוניים")
    add_code(
        doc,
        """
TaxExpense(i,t) = -TaxProvision(i,t)

RawTaxRate_Quarterly(i,t) = TaxExpense(i,t) / PretaxIncome(i,t)

NOPAT_Quarterly(i,t) = EBIT(i,t) * (1 - RawTaxRate_Quarterly(i,t))
""",
    )
    add_paragraph(
        doc,
        "GuruFocus מדווח tax_provision בסימן שלילי עבור הוצאת מס. "
        "כל רכיבי החישוב נלקחים מאותו רבעון; אין סכימה של ארבעה רבעונים.",
    )

    add_heading(doc, "2. IC_RAW")
    add_code(
        doc,
        """
IC_RAW = TCA - TCL + NetPPE + Goodwill
""",
    )
    add_paragraph(
        doc,
        "NetPPE הוא רכוש קבוע מוחשי נטו ולכן Goodwill אינו כלול בו. "
        "ב-IC_RAW המוניטין נוסף במפורש.",
    )

    add_heading(doc, "3. ממוצעי ההון שמתאימים למונה הרבעוני")
    add_code(
        doc,
        """
AverageIC_RAW_Quarterly = (IC_RAW(t-1) + IC_RAW(t)) / 2
""",
    )

    add_heading(doc, "4. ROIC רבעוני לפני מס ואחרי מס")
    add_code(
        doc,
        """
ROIC_Pretax_Quarterly_IC_RAW  = EBIT(i,t)  / AverageIC_RAW_Quarterly
ROIC_Posttax_Quarterly_IC_RAW = NOPAT(i,t) / AverageIC_RAW_Quarterly
""",
    )
    add_paragraph(
        doc,
        "היחסים מחושבים גם כאשר ההון הממוצע שלילי. רק מכנה אפס או חסר "
        "מוחזר כריק.",
    )

    add_heading(doc, "5. EV/FCF רבעוני")
    add_code(
        doc,
        """
EV = MarketCap + ShortTermDebtAndCapitalLease
     + LongTermDebtAndCapitalLease - Cash - STInvestments

FreeCashFlow_TTM = FreeCashFlow(i,t) + FreeCashFlow(i,t-1)
                 + FreeCashFlow(i,t-2) + FreeCashFlow(i,t-3)

EV_FCF_Quarterly = EV / FreeCashFlow_TTM
""",
    )
    add_paragraph(
        doc,
        "המונה הוא EV לסוף הרבעון והמכנה הוא סכום ה-FCF בארבעת הרבעונים "
        "הרצופים האחרונים, בהתאם למתודולוגיית GuruFocus.",
    )

    add_heading(doc, "6. ערך חוב ויחס חוב להון עצמי")
    add_code(
        doc,
        """
DebtValue = ShortTermDebtAndCapitalLease + LongTermDebtAndCapitalLease
Equity = TotalAssets - TotalLiabilities
DebtToEquity = DebtValue / TotalStockholdersEquity
""",
    )

    add_heading(doc, "7. פירוק Shapley של השינוי הרבעוני")
    add_code(
        doc,
        """
E = EBIT,  T = RawTaxRate,  Q = 1 - T,  I = AverageIC_RAW_Quarterly

NOPAT = E * Q          ROIC = E * Q / I

C_EBIT^NOPAT = (E1 - E0) * (Q0 + Q1) / 2
C_Tax^NOPAT  = (Q1 - Q0) * (E0 + E1) / 2

C_EBIT^ROIC = (E1-E0) * [ (1/3)(Q0/I0) + (1/6)(Q1/I0)
                        + (1/6)(Q0/I1) + (1/3)(Q1/I1) ]
C_Tax^ROIC  = (Q1-Q0) * [ (1/3)(E0/I0) + (1/6)(E1/I0)
                        + (1/6)(E0/I1) + (1/3)(E1/I1) ]
C_IC^ROIC   = (1/I1 - 1/I0) * [ (1/3)(E0*Q0) + (1/6)(E1*Q0)
                              + (1/6)(E0*Q1) + (1/3)(E1*Q1) ]
""",
    )
    add_paragraph(
        doc,
        "ערך Shapley הוא ממוצע התרומה השולית של כל גורם על פני כל סדרי "
        "השינוי האפשריים, ולכן הוא אינו תלוי בסדר וסכום התרומות שווה בדיוק "
        "לשינוי הכולל. פירוק סדרתי היה מייחס את איברי האינטראקציה לגורם "
        "שמשתנה אחרון.",
    )
    add_paragraph(
        doc,
        "המס נכנס כ-Q = 1 - T ולא כ-T, כדי ששני המדדים יישארו מכפלתיים גם "
        "כאשר שיעור המס שלילי או גבוה מ-100%. I נקרא מעמודת הממוצע הקיימת "
        "בשתי נקודות זמן ואינו ממוצע פעם נוספת.",
    )
    add_paragraph(
        doc,
        "סימן ההשפעה נקבע לפי תרומת Shapley ולא לפי כיוון התנועה הגולמי: "
        "כאשר Q שלילי עלייה ב-EBIT מקטינה את NOPAT, וכאשר NOPAT שלילי ירידה "
        "בהון המושקע הופכת את ROIC לשלילי יותר.",
    )

    add_heading(doc, "8. עמודות החישוב בפלט")
    add_table(
        doc,
        [
            ("calc_tax_expense_quarterly", "מספר", "הוצאת המס לרבעון לאחר היפוך סימן"),
            ("calc_raw_tax_rate_quarterly", "יחס", "TaxExpense / PretaxIncome באותו רבעון"),
            ("calc_nopat_quarterly", "מספר", "EBIT רבעוני כפול 1 פחות RawTaxRate"),
            ("calc_ic_raw", "מספר", "TCA - TCL + NetPPE + Goodwill"),
            ("calc_average_ic_raw_quarterly", "מספר", "ממוצע יתרת פתיחה וסגירה של IC_RAW"),
            ("calc_roic_pretax_quarterly_ic_raw", "יחס", "ROIC רבעוני לפני מס על IC_RAW"),
            ("calc_roic_posttax_quarterly_ic_raw", "יחס", "ROIC רבעוני אחרי מס על IC_RAW"),
            ("calc_debt_value_quarterly", "מספר", "חוב קצר וארוך כולל חכירות"),
            ("calc_debt_to_equity_quarterly", "יחס", "DebtValue / TotalStockholdersEquity"),
            ("calc_enterprise_value_quarterly", "מספר", "שווי פעילות לסוף הרבעון"),
            ("calc_free_cash_flow_ttm", "מספר", "סכום FCF בארבעת הרבעונים האחרונים"),
            ("calc_ev_to_fcf_quarterly", "מכפיל", "Quarter-end EV / trailing-four-quarter FCF"),
        ],
    )

    add_heading(doc, "9. עמודות הפירוק")
    add_table(doc, _decomposition_rows())
    add_paragraph(
        doc,
        "בנוסף נוצרות 63 עמודות דמה (9 + 27 + 27) במכפלה קרטזית, כך שכל "
        "שילוב אפשרי מקבל עמודה. בכל שורה מסווגת נדלקת בדיוק עמודה אחת מכל "
        "משפחה, ובשורה שאינה מסווגת כל העמודות מקבלות 0.",
    )

    add_heading(doc, "10. עלות ההון (WACC)")
    add_code(
        doc,
        """
WACC = E/(D+E) * Re  +  D/(D+E) * Rd * (1 - T)

E  = MarketCap                          שווי שוק, לא הון עצמי חשבונאי
D  = DebtValue                          חוב קצר וארוך כולל חכירות
Re = Rf + ERP                           ללא CAPM וללא בטא
Rd = InterestExpense_TTM / AverageDebt
     AverageDebt = (D(t-4) + D(t)) / 2
T  = TaxExpense_TTM / PretaxIncome_TTM  אפס אם המכנה אינו חיובי

ROIC_Annualized = (1 + ROIC_Quarterly) ^ 4     - 1
WACC_Quarterly  = (1 + WACC_Annual)    ^ (1/4) - 1
""",
    )
    add_paragraph(
        doc,
        "המטרה אינה הערכת שווי אלא אומדן אחיד לסינון. Rf ו-ERP נקבעים "
        "בקובץ ההגדרות ונכתבים כעמודות בכל שורה, כך שכל תצפית מתעדת את "
        "ההנחה שלפיה תומחרה.",
    )
    add_paragraph(
        doc,
        "המשקלים משתמשים ביתרת סוף התקופה, כמו שווי השוק. הממוצע קיים רק "
        "במכנה של עלות החוב, משום שהמונה שם הוא זרימה של שנה שלמה.",
    )
    add_paragraph(
        doc,
        "שתי ההמרות מוחזרות ריקות כאשר הבסיס אינו חיובי: בחזקה זוגית "
        "תשואה רבעונית של מינוס 200% הייתה יוצאת אפס, ובשורש רביעי מספר "
        "שלילי אינו מוגדר. מכיוון שההיוון מונוטוני, פסק הדין של ROIC מול "
        "WACC זהה בשתי היחידות.",
    )
    add_paragraph(
        doc,
        "ה-API מדווח הוצאות ריבית כאפס בחלק ניכר מהרבעונים גם כשקיים חוב. "
        "הערך נלקח כפי שהוא, ועמודת calc_wacc_quality_flag מסמנת את השורה "
        "כ-DEBT_WITHOUT_INTEREST כדי שניתן יהיה לסנן אותה.",
    )
    add_table(doc, _wacc_rows())

    return doc


def _wacc_rows() -> list[tuple[str, str, str]]:
    """שורות הטבלה של עמודות ה-WACC, לפי הסדר הקנוני שבמודול."""
    from gurufocus.wacc import wacc_columns

    meanings = {
        "calc_roic_pretax_annualized_ic_raw": ("יחס", "ROIC לפני מס בהיוון שנתי, (1+r)^4-1"),
        "calc_roic_posttax_annualized_ic_raw": ("יחס", "ROIC אחרי מס בהיוון שנתי, (1+r)^4-1"),
        "calc_wacc_risk_free_rate": ("יחס", "Rf מקובץ ההגדרות"),
        "calc_wacc_equity_risk_premium": ("יחס", "ERP מקובץ ההגדרות"),
        "calc_wacc_cost_of_equity": ("יחס", "Re = Rf + ERP"),
        "calc_wacc_equity_value": ("מספר", "E — שווי השוק"),
        "calc_wacc_average_debt": ("מספר", "(D(t-4)+D(t))/2, למכנה של עלות החוב בלבד"),
        "calc_wacc_total_capital": ("מספר", "E + D לסוף התקופה"),
        "calc_wacc_equity_weight": ("יחס", "E/(D+E)"),
        "calc_wacc_debt_weight": ("יחס", "D/(D+E)"),
        "calc_interest_expense_ttm": ("מספר", "הוצאות ריבית בארבעה רבעונים, בסימן חיובי"),
        "calc_wacc_cost_of_debt": ("יחס", "Rd — ריבית TTM חלקי חוב ממוצע"),
        "calc_wacc_tax_rate": ("יחס", "שיעור מס TTM עם רצפה באפס"),
        "calc_wacc_after_tax_cost_of_debt": ("יחס", "Rd * (1-T)"),
        "calc_wacc_annual": ("יחס", "עלות ההון המשוקללת, שיעור שנתי"),
        "calc_wacc_quarterly": ("יחס", "אותו שיעור בהמרה רבעונית גאומטרית"),
        "calc_wacc_quality_flag": ("תווית", "עד כמה ניתן לסמוך על עלות החוב"),
        "calc_wacc_inputs_complete": ("בוליאני", "האם כל הנתונים לחישוב היו זמינים"),
        "calc_roic_minus_wacc_annualized": ("יחס", "הפער בין ROIC ל-WACC, שנתי"),
        "calc_roic_minus_wacc_quarterly": ("יחס", "הפער בין ROIC ל-WACC, רבעוני"),
        "calc_creates_value": ("בוליאני", "ROIC גבוה מ-WACC, כלומר החברה יוצרת ערך"),
    }
    missing = [name for name in wacc_columns() if name not in meanings]
    if missing:
        raise RuntimeError("חסר תיאור לעמודות WACC: " + ", ".join(missing))
    return [(name, *meanings[name]) for name in wacc_columns()]


def _decomposition_rows() -> list[tuple[str, str, str]]:
    """שורות הטבלה של עמודות הפירוק, לפי הסדר הקנוני שבמודול."""
    from gurufocus.decomposition import NAMED_DECOMPOSITION_COLUMNS

    meanings = {
        "calc_nopat_decomposition_status": ("תווית", "האם ניתן להשוות לרבעון הקודם, ואם לא — מדוע"),
        "calc_nopat_change_quarterly": ("מספר", "NOPAT(t) פחות NOPAT(t-1)"),
        "calc_nopat_change_direction": ("תווית", "INCREASE / DECREASE / STABLE / UNCLASSIFIED"),
        "calc_nopat_ebit_contribution": ("מספר", "תרומת EBIT לשינוי ב-NOPAT"),
        "calc_nopat_tax_contribution": ("מספר", "תרומת שיעור המס לשינוי ב-NOPAT"),
        "calc_nopat_decomposition_residual": ("מספר", "שארית בדיקה, אפס עד כדי דיוק מספרי"),
        "calc_nopat_ebit_effect": ("תווית", "סימן תרומת EBIT: POSITIVE / NEGATIVE / NEUTRAL"),
        "calc_nopat_tax_effect": ("תווית", "סימן תרומת המס"),
        "calc_nopat_effect_combination": ("תווית", "שילוב שני הסימנים, אחד מתשעה"),
        "calc_roic_decomposition_status": ("תווית", "כמו למעלה, בתוספת UNCLASSIFIED_ZERO_IC"),
        "calc_roic_posttax_change_quarterly": ("יחס", "ROIC(t) פחות ROIC(t-1)"),
        "calc_roic_posttax_change_direction": ("תווית", "כיוון השינוי ב-ROIC"),
        "calc_roic_ebit_contribution": ("יחס", "תרומת EBIT לשינוי ב-ROIC"),
        "calc_roic_tax_contribution": ("יחס", "תרומת שיעור המס לשינוי ב-ROIC"),
        "calc_roic_ic_contribution": ("יחס", "תרומת ההון המושקע לשינוי ב-ROIC"),
        "calc_roic_decomposition_residual": ("יחס", "שארית בדיקה, אפס עד כדי דיוק מספרי"),
        "calc_roic_ebit_effect": ("תווית", "סימן תרומת EBIT ל-ROIC"),
        "calc_roic_tax_effect": ("תווית", "סימן תרומת המס ל-ROIC"),
        "calc_roic_ic_effect": ("תווית", "סימן תרומת ההון המושקע ל-ROIC"),
        "calc_roic_effect_combination": ("תווית", "שילוב שלושת הסימנים, אחד מ-27"),
        "calc_ebit_movement": ("תווית", "כיוון גולמי של EBIT: UP / DOWN / FLAT"),
        "calc_tax_rate_movement": ("תווית", "כיוון גולמי של שיעור המס"),
        "calc_ic_movement": ("תווית", "כיוון גולמי של ההון המושקע"),
        "calc_raw_movement_combination": ("תווית", "שילוב שלוש התנועות הגולמיות, אחד מ-27"),
        "calc_roic_has_opposing_effects": ("בוליאני", "קיימת תרומה חיובית וגם תרומה שלילית"),
        "calc_roic_positive_driver_count": ("מספר", "כמה תרומות חיוביות"),
        "calc_roic_negative_driver_count": ("מספר", "כמה תרומות שליליות"),
        "calc_roic_neutral_driver_count": ("מספר", "כמה תרומות בתוך סף המהותיות"),
        "calc_roic_active_driver_count": ("מספר", "כמה תרומות אינן ניטרליות"),
        "calc_roic_effect_structure": ("תווית", "מבנה התרומות: ALL_POSITIVE, MIXED_FULL_OFFSET וכו'"),
        "calc_roic_total_absolute_contribution": ("יחס", "סכום הערכים המוחלטים של שלוש התרומות"),
        "calc_roic_ebit_absolute_share": ("יחס", "משקל התרומה המוחלטת של EBIT"),
        "calc_roic_tax_absolute_share": ("יחס", "משקל התרומה המוחלטת של המס"),
        "calc_roic_ic_absolute_share": ("יחס", "משקל התרומה המוחלטת של ההון המושקע"),
        "calc_roic_dominant_driver": ("תווית", "EBIT / TAX / IC / BALANCED / NONE"),
        "calc_roic_dominant_driver_effect": ("תווית", "סימן התרומה של הגורם הדומיננטי"),
        "calc_roic_offset_ratio": ("יחס", "1 פחות |ΔROIC| חלקי סכום התרומות המוחלטות"),
        "calc_ebit_sign_regime": ("תווית", "מעבר רווח/הפסד/אפס של EBIT"),
        "calc_nopat_sign_regime": ("תווית", "מעבר רווח/הפסד/אפס של NOPAT"),
        "calc_tax_rate_quality_flag": ("תווית", "VALID / NEGATIVE_TAX_RATE / ABOVE_100_PERCENT / MISSING"),
        "calc_roic_quality_flag": ("תווית", "מסמן הון מושקע שלילי או שיעור מס מחוץ לטווח"),
        "calc_roic_economic_interpretation_valid": ("בוליאני", "האם ניתן לקרוא את היחס כיעילות כלכלית"),
        "calc_roic_business_classification": ("תווית", "סיווג עסקי מסכם, ממצה לכל שילוב"),
        "calc_roic_explanation": ("טקסט", "משפט קצר באנגלית הנגזר מהתרומות"),
    }
    missing = [name for name in NAMED_DECOMPOSITION_COLUMNS if name not in meanings]
    if missing:
        raise RuntimeError("חסר תיאור לעמודות פירוק: " + ", ".join(missing))
    return [(name, *meanings[name]) for name in NAMED_DECOMPOSITION_COLUMNS]


def main() -> int:
    document = build_document()
    document.save(OUT_PATH)
    print(f"נשמר: {OUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
